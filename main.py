import os
import json
import re
import time
import uuid
from fastapi import FastAPI, HTTPException, UploadFile, File, Form, Header, Depends
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel, EmailStr
from huggingface_hub import InferenceClient
from typing import Optional, List
from supabase import create_client, Client

app = FastAPI()
app.add_middleware(CORSMiddleware, allow_origins=["*"], allow_methods=["*"], allow_headers=["*"])

HF_TOKEN = os.environ.get("HF_TOKEN", "")

# ---------------------------------------------------------------------------
# Supabase setup (Auth + Postgres + Storage)
# ---------------------------------------------------------------------------
SUPABASE_URL = os.environ.get("SUPABASE_URL", "")
SUPABASE_SERVICE_KEY = os.environ.get("SUPABASE_SERVICE_KEY", "")
TEACHER_INVITE_CODE = os.environ.get("TEACHER_INVITE_CODE", "")
DOCUMENTS_BUCKET = "documents"

supabase: Optional[Client] = (
    create_client(SUPABASE_URL, SUPABASE_SERVICE_KEY)
    if SUPABASE_URL and SUPABASE_SERVICE_KEY
    else None
)

ALLOWED_CATEGORIES = {"ielts", "hsca", "thpt"}
ALLOWED_ROLES = {"student", "teacher"}


def require_supabase():
    if not supabase:
        raise HTTPException(status_code=500, detail="Auth/Database chưa được cấu hình (thiếu SUPABASE_URL/SUPABASE_SERVICE_KEY)")


def get_current_user(authorization: Optional[str] = Header(None)):
    require_supabase()
    if not authorization or not authorization.startswith("Bearer "):
        raise HTTPException(status_code=401, detail="Thiếu token đăng nhập")
    token = authorization.split(" ", 1)[1]
    try:
        user_resp = supabase.auth.get_user(token)
        user = user_resp.user
        if not user:
            raise HTTPException(status_code=401, detail="Token không hợp lệ")
    except HTTPException:
        raise
    except Exception:
        raise HTTPException(status_code=401, detail="Token không hợp lệ hoặc đã hết hạn")

    profile = supabase.table("profiles").select("*").eq("id", user.id).single().execute()
    if not profile.data:
        raise HTTPException(status_code=401, detail="Không tìm thấy hồ sơ người dùng")
    return {"id": user.id, "email": user.email, **profile.data}


# ---------------------------------------------------------------------------
# Auth
# ---------------------------------------------------------------------------
class RegisterRequest(BaseModel):
    email: EmailStr
    password: str
    full_name: str
    role: str  # "student" | "teacher"
    invite_code: Optional[str] = None


class LoginRequest(BaseModel):
    email: EmailStr
    password: str


@app.post("/auth/register")
def register(req: RegisterRequest):
    require_supabase()
    if req.role not in ALLOWED_ROLES:
        raise HTTPException(status_code=400, detail="Vai trò không hợp lệ")
    if len(req.password) < 6:
        raise HTTPException(status_code=400, detail="Mật khẩu phải có ít nhất 6 ký tự")

    if req.role == "teacher":
        if not TEACHER_INVITE_CODE:
            raise HTTPException(status_code=500, detail="Hệ thống chưa cấu hình mã mời giáo viên")
        if req.invite_code != TEACHER_INVITE_CODE:
            raise HTTPException(status_code=403, detail="Mã mời giáo viên không đúng")

    try:
        auth_resp = supabase.auth.admin.create_user({
            "email": req.email,
            "password": req.password,
            "email_confirm": True,
        })
        user_id = auth_resp.user.id
        supabase.table("profiles").insert({
            "id": user_id,
            "full_name": req.full_name,
            "role": req.role,
        }).execute()
        return {"status": "ok", "message": "Đăng ký thành công, mời bạn đăng nhập"}
    except Exception as e:
        msg = str(e)
        if "already" in msg.lower() or "duplicate" in msg.lower():
            msg = "Email này đã được đăng ký"
        raise HTTPException(status_code=400, detail=msg)


@app.post("/auth/login")
def login(req: LoginRequest):
    require_supabase()
    try:
        resp = supabase.auth.sign_in_with_password({"email": req.email, "password": req.password})
        profile = supabase.table("profiles").select("*").eq("id", resp.user.id).single().execute()
        return {
            "access_token": resp.session.access_token,
            "user": {"id": resp.user.id, "email": resp.user.email, **(profile.data or {})},
        }
    except Exception:
        raise HTTPException(status_code=401, detail="Email hoặc mật khẩu không đúng")


@app.get("/me")
def me(user=Depends(get_current_user)):
    return user


# ---------------------------------------------------------------------------
# Documents (IELTS / HSCA / THPT libraries)
# ---------------------------------------------------------------------------
@app.get("/documents")
def list_documents(category: str, subject: Optional[str] = None, grade: Optional[str] = None):
    require_supabase()
    if category not in ALLOWED_CATEGORIES:
        raise HTTPException(status_code=400, detail="Danh mục không hợp lệ")
    q = supabase.table("documents").select("*").eq("category", category).order("created_at", desc=True)
    if subject:
        q = q.eq("subject", subject)
    if grade:
        q = q.eq("grade", grade)
    return q.execute().data


@app.post("/documents/upload")
async def upload_document(
    category: str = Form(...),
    subject: str = Form(""),
    grade: str = Form(""),
    title: str = Form(...),
    file: UploadFile = File(...),
    user=Depends(get_current_user),
):
    if user["role"] != "teacher":
        raise HTTPException(status_code=403, detail="Chỉ giáo viên mới được đăng tài liệu")
    if category not in ALLOWED_CATEGORIES:
        raise HTTPException(status_code=400, detail="Danh mục không hợp lệ")

    file_bytes = await file.read()
    if len(file_bytes) > 20 * 1024 * 1024:
        raise HTTPException(status_code=400, detail="File vượt quá 20MB")

    ext = os.path.splitext(file.filename or "")[1]
    storage_path = f"{category}/{uuid.uuid4().hex}{ext}"

    try:
        supabase.storage.from_(DOCUMENTS_BUCKET).upload(
            storage_path,
            file_bytes,
            {"content-type": file.content_type or "application/octet-stream"},
        )
        public_url = supabase.storage.from_(DOCUMENTS_BUCKET).get_public_url(storage_path)

        row = {
            "title": title,
            "category": category,
            "subject": subject,
            "grade": grade,
            "file_url": public_url,
            "file_name": file.filename,
            "uploaded_by": user["id"],
            "uploader_name": user.get("full_name", ""),
        }
        supabase.table("documents").insert(row).execute()
        return {"status": "ok", "file_url": public_url}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.delete("/documents/{doc_id}")
def delete_document(doc_id: str, user=Depends(get_current_user)):
    if user["role"] != "teacher":
        raise HTTPException(status_code=403, detail="Chỉ giáo viên mới được xoá tài liệu")
    doc = supabase.table("documents").select("*").eq("id", doc_id).single().execute()
    if not doc.data:
        raise HTTPException(status_code=404, detail="Không tìm thấy tài liệu")
    if doc.data["uploaded_by"] != user["id"]:
        raise HTTPException(status_code=403, detail="Bạn chỉ có thể xoá tài liệu do chính mình đăng")
    supabase.table("documents").delete().eq("id", doc_id).execute()
    return {"status": "ok"}


# ---------------------------------------------------------------------------
# IELTS Writing grader (unchanged from existing production logic)
# ---------------------------------------------------------------------------
TASK1_DESCRIPTORS = """
IELTS WRITING TASK 1 - OFFICIAL BAND DESCRIPTORS (Updated May 2023)
Each criterion is scored as a WHOLE NUMBER only (1-9). No decimals. No .5.

TASK ACHIEVEMENT (TA):
9: All requirements fully and appropriately satisfied. Extremely rare lapses.
8: Covers all requirements appropriately, relevantly, sufficiently. Key features skilfully selected, presented, highlighted, illustrated. Occasional omissions.
7: Covers requirements. Content relevant and accurate with few omissions. Key features covered and clearly highlighted. Clear overview, data appropriately categorised, main trends identified.
6: Focuses on requirements. Key features covered and adequately highlighted. Overview attempted. Some irrelevant/inaccurate info may occur. Some details missing or excessive.
5: Generally addresses requirements. Key features not adequately covered. Recounting mainly mechanical. May be no data to support description. Tendency to focus on details without bigger picture.
4: Attempts to address task. Few key features selected. Key features may be irrelevant, repetitive, inaccurate. Format may be inappropriate.
3: Does not address requirements (possibly misunderstood data/diagram). Key features largely irrelevant. Limited information, used repetitively.
2: Content barely relates to task.
1: Content wholly unrelated to task (or 20 words or fewer).

COHERENCE & COHESION (CC):
9: Message followed effortlessly. Cohesion rarely attracts attention. Minimal lapses. Paragraphing skilfully managed.
8: Message followed with ease. Ideas logically sequenced, cohesion well managed. Occasional lapses. Paragraphing used sufficiently and appropriately.
7: Ideas logically organised, clear progression. Few lapses. Cohesive devices used flexibly but with some inaccuracies or over/under use.
6: Generally arranged coherently, clear overall progression. Cohesive devices used to some good effect but cohesion may be faulty or mechanical. Reference/substitution may lack flexibility.
5: Organisation evident but not wholly logical, may lack overall progression. Relationship of ideas can be followed but sentences not fluently linked. Limited/overuse of cohesive devices.
4: Ideas evident but not arranged coherently, no clear progression. Relationships unclear/inadequately marked. Inaccurate use or lack of substitution/referencing.
3: No apparent logical organisation. Minimal use of sequencers/cohesive devices. Difficulty identifying referencing.
2: Little relevant message. Little evidence of control of organisational features.
1: Writing fails to communicate any message.

LEXICAL RESOURCE (LR):
9: Full flexibility and precise use evident. Wide range used accurately and appropriately. Very natural and sophisticated control. Minor errors extremely rare.
8: Wide resource fluently and flexibly used to convey precise meanings. Skilful use of uncommon/idiomatic items. Occasional errors in spelling/word formation with minimal impact.
7: Sufficient for some flexibility and precision. Some ability to use less common/idiomatic items. Awareness of style and collocation, though inappropriacies occur. Few errors in spelling/word formation.
6: Generally adequate and appropriate. Meaning generally clear despite restricted range or lack of precision. Some errors in spelling/word formation but do not impede communication.
5: Limited but minimally adequate. Simple vocabulary used accurately but range doesn't permit much variation. Frequent lapses in appropriacy. Errors in spelling/word formation may cause difficulty.
4: Limited and inadequate for task. Basic vocabulary, may be used repetitively. Inappropriate word choice/errors in word formation may impede meaning.
3: Inadequate. Possible over-dependence on memorised language. Control of word choice/spelling very limited, errors predominate and may severely impede meaning.
2: Extremely limited. Few recognisable strings apart from memorised phrases. No apparent control of word formation/spelling.
1: No resource apparent except a few isolated words.

GRAMMATICAL RANGE & ACCURACY (GRA):
9: Wide range of structures used with full flexibility and control. Punctuation and grammar used appropriately throughout. Minor errors extremely rare.
8: Wide range of structures flexibly and accurately used. Majority of sentences error-free, punctuation well managed. Occasional non-systematic errors with minimal impact.
7: Variety of complex structures used with some flexibility and accuracy. Grammar and punctuation generally well controlled. Error-free sentences frequent. Few errors persist but don't impede communication.
6: Mix of simple and complex sentence forms but flexibility limited. More complex structures not as accurate as simple. Errors in grammar and punctuation occur but rarely impede communication.
5: Range of structures limited and rather repetitive. Complex sentences attempted but tend to be faulty. Grammatical errors may be frequent and cause difficulty. Punctuation may be faulty.
4: Very limited range of structures. Subordinate clauses rare, simple sentences predominate. Grammatical errors frequent and may impede meaning. Punctuation often faulty.
3: Sentence forms attempted but errors in grammar and punctuation predominate. Prevents most meaning from coming through. Length may be insufficient.
2: Little or no evidence of sentence forms except in memorised phrases.
1: No rateable language evident.

OVERALL BAND CALCULATION for Task 1:
Average = (TA + CC + LR + GRA) / 4
Round to nearest 0.5: if average ends in .25 round up to .5, if ends in .75 round up to next whole number.
Example: (6+6+7+6)/4 = 6.25 → Overall Band 6.5
"""

TASK2_DESCRIPTORS = """
IELTS WRITING TASK 2 - OFFICIAL BAND DESCRIPTORS (Updated May 2023)
Each criterion is scored as a WHOLE NUMBER only (1-9). No decimals. No .5.

TASK RESPONSE (TR):
9: Prompt appropriately addressed and explored in depth. Clear and fully developed position which directly answers question. Ideas relevant, fully extended and well supported.
8: Prompt appropriately and sufficiently addressed. Clear and well-developed position. Ideas relevant, well extended and supported. Occasional omissions.
7: Main parts of prompt appropriately addressed. Clear and developed position. Main ideas extended and supported but may over-generalise or lack focus/precision.
6: Main parts addressed (some more fully than others). Position directly relevant to prompt but conclusions may be unclear/unjustified. Main ideas relevant but some insufficiently developed.
5: Main parts incompletely addressed. Format may be inappropriate. Position expressed but development not always clear. Some main ideas put forward but limited and not sufficiently developed.
4: Prompt tackled minimally or tangentially. Position discernible but hard to find. Main ideas difficult to identify. Large parts may be repetitive.
3: No part adequately addressed or prompt misunderstood. No relevant position. Few ideas, may be irrelevant or insufficiently developed.
2: Content barely related to prompt. No position identifiable.
1: Content wholly unrelated to prompt (or 20 words or fewer).

COHERENCE & COHESION (CC):
9: Message followed effortlessly. Cohesion rarely attracts attention. Minimal lapses. Paragraphing skilfully managed.
8: Message followed with ease. Ideas logically sequenced, cohesion well managed. Occasional lapses. Paragraphing used sufficiently and appropriately.
7: Ideas logically organised, clear progression throughout. Few minor lapses. Cohesive devices used flexibly but with some inaccuracies or over/under use. Paragraphing generally effective.
6: Generally arranged coherently, clear overall progression. Cohesive devices used to some good effect but cohesion may be faulty or mechanical. Paragraphing may not always be logical.
5: Organisation evident but not wholly logical, may lack overall progression. Relationship of ideas can be followed but sentences not fluently linked. Paragraphing may be inadequate or missing.
4: Ideas evident but not arranged coherently, no clear progression. Relationships unclear. May be no paragraphing or no clear main topic within paragraphs.
3: No apparent logical organisation. Minimal use of sequencers/cohesive devices. Difficulty identifying referencing. Any paragraphing attempts are unhelpful.
2: Little relevant message. Little evidence of control of organisational features.
1: Writing fails to communicate any message.

LEXICAL RESOURCE (LR):
9: Full flexibility and precise use widely evident. Wide range used accurately and appropriately. Very natural and sophisticated control. Minor errors extremely rare.
8: Wide resource fluently and flexibly used to convey precise meanings. Skilful use of uncommon/idiomatic items. Occasional errors with minimal impact.
7: Sufficient for some flexibility and precision. Some ability to use less common/idiomatic items. Awareness of style and collocation, though inappropriacies occur. Few errors.
6: Generally adequate and appropriate. Meaning generally clear despite restricted range or lack of precision. Some errors but do not impede communication.
5: Limited but minimally adequate. Simple vocabulary used accurately but range doesn't permit much variation. Frequent lapses in appropriacy. Errors may cause difficulty.
4: Limited and inadequate for task. Basic vocabulary, may be used repetitively. Errors in word formation/spelling may impede meaning.
3: Inadequate. Possible over-dependence on memorised language. Control very limited, errors predominate and may severely impede meaning.
2: Extremely limited. Few recognisable strings. No apparent control.
1: No resource apparent.

GRAMMATICAL RANGE & ACCURACY (GRA):
9: Wide range of structures used with full flexibility and control. Grammar and punctuation used appropriately throughout. Minor errors extremely rare.
8: Wide range flexibly and accurately used. Majority of sentences error-free. Occasional non-systematic errors with minimal impact.
7: Variety of complex structures with some flexibility and accuracy. Grammar and punctuation generally well controlled. Error-free sentences frequent. Few errors persist.
6: Mix of simple and complex forms but flexibility limited. More complex structures not as accurate. Errors in grammar and punctuation occur but rarely impede communication.
5: Range limited and rather repetitive. Complex sentences tend to be faulty. Grammatical errors may be frequent and cause difficulty. Punctuation may be faulty.
4: Very limited range. Simple sentences predominate. Grammatical errors frequent and may impede meaning. Punctuation often faulty.
3: Sentence forms attempted but errors predominate. Prevents most meaning from coming through.
2: Little or no evidence of sentence forms.
1: No rateable language evident.

OVERALL BAND CALCULATION for Task 2:
Average = (TR + CC + LR + GRA) / 4
Round to nearest 0.5: if average ends in .25 round up to .5, if ends in .75 round up to next whole number.
Example: (6+7+6+6)/4 = 6.25 → Overall Band 6.5
Example: (7+7+7+6)/4 = 6.75 → Overall Band 7.0
"""


class GradeRequest(BaseModel):
    essay: str
    prompt: str
    task_type: str
    image_base64: Optional[str] = None


@app.get("/")
def root():
    return {"status": "LND Academy API is running"}


@app.post("/grade")
def grade_essay(req: GradeRequest):
    if not HF_TOKEN:
        raise HTTPException(status_code=500, detail="HF_TOKEN not configured")
    if len(req.essay.strip()) < 50:
        raise HTTPException(status_code=400, detail="Essay is too short")

    client = InferenceClient(provider="nscale", api_key=HF_TOKEN)

    if req.task_type == "Task 1":
        criterion1_name = "Task Achievement (TA)"
        descriptors = TASK1_DESCRIPTORS
    else:
        criterion1_name = "Task Response (TR)"
        descriptors = TASK2_DESCRIPTORS

    system_prompt = (
        "You are a certified IELTS examiner with 10+ years of experience. "
        "You must grade strictly and accurately following the official IELTS Band Descriptors below.\n\n"
        + descriptors +
        "\n\nCRITICAL RULES:\n"
        "1. Each criterion score MUST be a WHOLE NUMBER (1-9). NEVER use decimals like 6.5 for individual criteria.\n"
        "2. Only the overall_band can be x.0 or x.5 (calculated by averaging the 4 criteria and rounding to nearest 0.5).\n"
        "3. Read the band descriptors carefully for EACH criterion before scoring.\n"
        "4. Be strict and accurate — do not inflate scores.\n"
        "5. Respond ONLY in this exact JSON format, no extra text, no markdown:\n"
        "{\n"
        f'  "criterion1_name": "{criterion1_name}",\n'
        '  "criterion1_score": 6,\n'
        '  "criterion1_feedback": "Specific feedback referencing the band descriptors...",\n'
        '  "criterion2_name": "Coherence & Cohesion",\n'
        '  "criterion2_score": 6,\n'
        '  "criterion2_feedback": "Specific feedback referencing the band descriptors...",\n'
        '  "criterion3_name": "Lexical Resource",\n'
        '  "criterion3_score": 7,\n'
        '  "criterion3_feedback": "Specific feedback referencing the band descriptors...",\n'
        '  "criterion4_name": "Grammatical Range & Accuracy",\n'
        '  "criterion4_score": 6,\n'
        '  "criterion4_feedback": "Specific feedback referencing the band descriptors...",\n'
        '  "overall_band": 6.5,\n'
        '  "overall_feedback": "Comprehensive examiner feedback on the essay overall..."\n'
        "}"
    )

    user_content = f"EXAM PROMPT:\n{req.prompt}\n\nSTUDENT ESSAY:\n{req.essay}"
    user_message = f"Grade this IELTS Writing {req.task_type} using the official band descriptors provided:\n\n{user_content}\n/no_think"

    messages = [
        {"role": "system", "content": system_prompt},
        {"role": "user", "content": user_message}
    ]

    try:
        response = client.chat.completions.create(
            model="Qwen/Qwen3-32B",
            messages=messages,
            max_tokens=2000,
        )
        raw = response.choices[0].message.content.strip()
        raw = re.sub(r'<think>.*?</think>', '', raw, flags=re.DOTALL).strip()
        if "```" in raw:
            raw = raw.split("```")[1]
            if raw.startswith("json"):
                raw = raw[4:]
        match = re.search(r'\{.*\}', raw, re.DOTALL)
        if match:
            raw = match.group(0)
        result = json.loads(raw.strip())

        for key in ["criterion1_score", "criterion2_score", "criterion3_score", "criterion4_score"]:
            if key in result:
                result[key] = int(round(result[key]))

        scores = [result["criterion1_score"], result["criterion2_score"],
                  result["criterion3_score"], result["criterion4_score"]]
        avg = sum(scores) / 4
        result["overall_band"] = round(avg * 2) / 2

        return result
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


# ---------------------------------------------------------------------------
# Live chat assistant (site-wide widget)
# ---------------------------------------------------------------------------
CHAT_SYSTEM_PROMPT = """Bạn là trợ lý ảo của LND Academy — trung tâm tiếng Anh học thuật chất lượng cao tại tỉnh Lâm Đồng.

VỀ TRUNG TÂM:
- Sứ mệnh: đồng hành và nâng tầm năng lực Anh ngữ học thuật cho học sinh Lâm Đồng, biến tiếng Anh thành môn học được yêu thích chứ không phải nỗi sợ, giúp học sinh thật sự dùng được ngôn ngữ để giao tiếp chứ không chỉ học để lấy điểm.
- Triết lý đào tạo: "Trust the process. See your progress." — kiên trì theo lộ trình khoa học, cá nhân hóa.

LỘ TRÌNH ĐÀO TẠO IELTS:
- Foundation: đầu vào 3.0 → đầu ra 5.5. Xây nền ngữ pháp, phát âm cốt lõi.
- Intermediate: đầu vào 5.5 → đầu ra 6.5. Phát triển tư duy nghị luận, triển khai đều 4 kỹ năng.
- Advanced: đầu vào 6.5 → đầu ra 7.0+. Tối ưu chiến thuật làm bài, hướng điểm số xuất sắc.
- Đặc điểm chung: sĩ số nhỏ cực hạn, tương tác sâu sát, sửa bài 1:1 chuyên sâu.

CHƯƠNG TRÌNH H-SCA (ôn thi Đánh giá năng lực chuyên biệt của Đại học Sư phạm TP.HCM):
- H-SCA Prep Course: lộ trình dài hơi, hệ thống hóa kiến thức trọng tâm ngữ pháp & từ vựng.
- H-SCA Crash Course: tổng ôn và luyện đề cấp tốc trong thời gian ngắn trước kỳ thi.

TÍNH NĂNG TRÊN WEBSITE:
- "Chấm bài Writing": chấm điểm IELTS Writing Task 1/Task 2 tự động bằng AI theo Band Descriptors chính thức, có nhận xét chi tiết 4 tiêu chí — dùng miễn phí, không cần đăng nhập.
- Thư viện tài liệu IELTS / HSCA / THPT: học sinh xem và tải miễn phí; giáo viên (đã xác minh qua mã mời) mới đăng tài liệu mới được.
- Cần đăng ký tài khoản để giáo viên đăng tài liệu; học sinh có thể xem tài liệu mà không cần đăng nhập.

ĐỊNH HƯỚNG SẮP RA MẮT: Ôn luyện Chuyên Anh (thi vào THPT chuyên), Bồi dưỡng Học sinh Giỏi cấp Tỉnh/Khu vực.

LIÊN HỆ:
- Hotline/Zalo: 0389 339 171 (Võ Sỹ Đồng) hoặc 0783 630 468 (Trần Long Nguyên)
- Email: lndacademy.work@gmail.com

QUY TẮC TRẢ LỜI:
- Trả lời tiếng Việt, thân thiện, ngắn gọn (2-5 câu), đi thẳng trọng tâm câu hỏi.
- Với câu hỏi học thuật (ngữ pháp, từ vựng, cấu trúc bài thi IELTS...), giải thích rõ ràng, dễ hiểu, có ví dụ ngắn nếu cần.
- Với câu hỏi về tuyển sinh (học phí cụ thể, lịch khai giảng, ưu đãi...) mà không có trong dữ liệu trên, KHÔNG bịa số liệu — hướng dẫn liên hệ hotline/Zalo/email để được tư vấn chính xác.
- Không tự ý cam kết học phí, ưu đãi, cam kết đầu ra cụ thể nếu không có trong dữ liệu trên."""


class ChatMessage(BaseModel):
    role: str
    content: str


class ChatRequest(BaseModel):
    messages: list[ChatMessage]


@app.post("/chat")
def chat(req: ChatRequest):
    if not HF_TOKEN:
        raise HTTPException(status_code=500, detail="HF_TOKEN not configured")
    if not req.messages:
        raise HTTPException(status_code=400, detail="Thiếu tin nhắn")

    client = InferenceClient(provider="nscale", api_key=HF_TOKEN)

    history = [{"role": m.role, "content": m.content} for m in req.messages[-12:]]
    messages = [{"role": "system", "content": CHAT_SYSTEM_PROMPT}] + history
    if messages[-1]["role"] == "user":
        messages[-1]["content"] += "\n/no_think"

    try:
        response = client.chat.completions.create(
            model="Qwen/Qwen3-32B",
            messages=messages,
            max_tokens=600,
        )
        raw = response.choices[0].message.content.strip()
        raw = re.sub(r'<think>.*?</think>', '', raw, flags=re.DOTALL).strip()
        return {"reply": raw}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


# ---------------------------------------------------------------------------
# Content bot + QC bot (grounded in the marketing masterplan on Google Sheets)
# ---------------------------------------------------------------------------
from google.oauth2 import service_account
from googleapiclient.discovery import build

GOOGLE_SERVICE_ACCOUNT_JSON = os.environ.get("GOOGLE_SERVICE_ACCOUNT_JSON", "")
MASTERPLAN_SHEET_ID = os.environ.get("MASTERPLAN_SHEET_ID", "")
MASTERPLAN_CACHE_TTL = 600  # giay - tranh goi Google Sheets API lien tuc

_masterplan_cache = {"text": None, "fetched_at": 0.0}


def _get_sheets_service():
    if not GOOGLE_SERVICE_ACCOUNT_JSON or not MASTERPLAN_SHEET_ID:
        raise HTTPException(
            status_code=500,
            detail="Chua cau hinh GOOGLE_SERVICE_ACCOUNT_JSON / MASTERPLAN_SHEET_ID",
        )
    try:
        info = json.loads(GOOGLE_SERVICE_ACCOUNT_JSON)
    except json.JSONDecodeError:
        raise HTTPException(status_code=500, detail="GOOGLE_SERVICE_ACCOUNT_JSON khong phai JSON hop le")
    creds = service_account.Credentials.from_service_account_info(
        info, scopes=["https://www.googleapis.com/auth/spreadsheets.readonly"]
    )
    return build("sheets", "v4", credentials=creds)


def _fetch_masterplan_text(max_chars: int = 6000) -> str:
    service = _get_sheets_service()
    sheet = service.spreadsheets()
    try:
        meta = sheet.get(spreadsheetId=MASTERPLAN_SHEET_ID).execute()
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Khong doc duoc Google Sheet: {e}")

    lines = []
    for s in meta.get("sheets", []):
        title = s["properties"]["title"]
        try:
            result = sheet.values().get(spreadsheetId=MASTERPLAN_SHEET_ID, range=title).execute()
        except Exception:
            continue
        values = result.get("values", [])
        if values:
            lines.append(f"--- Sheet: {title} ---")
        for row in values:
            line = " | ".join(str(c) for c in row if str(c).strip())
            if line.strip():
                lines.append(line)

    text = "\n".join(lines)
    return text[:max_chars]


def get_masterplan_text() -> str:
    now = time.time()
    if _masterplan_cache["text"] is None or (now - _masterplan_cache["fetched_at"]) > MASTERPLAN_CACHE_TTL:
        _masterplan_cache["text"] = _fetch_masterplan_text()
        _masterplan_cache["fetched_at"] = now
    return _masterplan_cache["text"]


PLATFORM_RULES = {
    "facebook": "Viết bài đăng Facebook/Zalo: khoảng 80-150 từ, giọng gần gũi, có thể dùng emoji vừa phải, kết thúc bằng lời kêu gọi hành động rõ ràng (liên hệ hotline/Zalo).",
    "blog": "Viết bài blog cho website: khoảng 400-700 từ, có tiêu đề, chia đoạn/heading phụ rõ ràng, giọng chuyên nghiệp nhưng dễ đọc.",
}


class ContentGenerateRequest(BaseModel):
    platform: str  # "facebook" | "blog"
    topic: str
    notes: Optional[str] = None


@app.post("/content/generate")
def generate_content(req: ContentGenerateRequest, user=Depends(get_current_user)):
    if user["role"] != "teacher":
        raise HTTPException(status_code=403, detail="Chỉ giáo viên/quản trị mới dùng được công cụ này")
    if req.platform not in PLATFORM_RULES:
        raise HTTPException(status_code=400, detail="Nền tảng không hợp lệ")
    if not HF_TOKEN:
        raise HTTPException(status_code=500, detail="HF_TOKEN not configured")

    masterplan = get_masterplan_text()

    system_prompt = f"""Bạn là chuyên viên content marketing của LND Academy — trung tâm tiếng Anh học thuật tại Lâm Đồng.
Dựa vào kế hoạch truyền thông (masterplan) trích từ Google Sheets dưới đây, viết nội dung theo đúng định hướng, chủ đề chiến dịch và giọng văn thương hiệu hiện tại. KHÔNG bịa số liệu/ưu đãi/học phí không có trong masterplan hoặc không được cung cấp.

KẾ HOẠCH TRUYỀN THÔNG:
{masterplan}

YÊU CẦU ĐỊNH DẠNG: {PLATFORM_RULES[req.platform]}
"""
    user_msg = f"Chủ đề/brief: {req.topic}"
    if req.notes:
        user_msg += f"\nGhi chú thêm: {req.notes}"
    user_msg += "\n/no_think"

    client = InferenceClient(provider="nscale", api_key=HF_TOKEN)
    try:
        response = client.chat.completions.create(
            model="Qwen/Qwen3-32B",
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_msg},
            ],
            max_tokens=1200,
        )
        raw = response.choices[0].message.content.strip()
        raw = re.sub(r'<think>.*?</think>', '', raw, flags=re.DOTALL).strip()
        return {"draft": raw}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


class ContentQCRequest(BaseModel):
    platform: str
    content: str


@app.post("/content/qc")
def qc_content(req: ContentQCRequest, user=Depends(get_current_user)):
    if user["role"] != "teacher":
        raise HTTPException(status_code=403, detail="Chỉ giáo viên/quản trị mới dùng được công cụ này")
    if not HF_TOKEN:
        raise HTTPException(status_code=500, detail="HF_TOKEN not configured")

    masterplan = get_masterplan_text()
    length_rule = {
        "facebook": "80-150 từ cho bài Facebook/Zalo",
        "blog": "400-700 từ cho bài blog",
    }.get(req.platform, "độ dài phù hợp nền tảng")

    system_prompt = f"""Bạn là chuyên viên Quality Control (QC) nội dung marketing của LND Academy.
Kiểm tra bài viết dưới đây theo các tiêu chí:
1. Giọng văn có nhất quán với thương hiệu LND Academy không (đối chiếu với masterplan bên dưới).
2. Có thông tin nào bị bịa đặt (học phí, ưu đãi, cam kết đầu ra) không nằm trong masterplan/dữ liệu đã biết không.
3. Chính tả và ngữ pháp tiếng Việt có lỗi không (liệt kê cụ thể nếu có).
4. Độ dài có phù hợp không: {length_rule}.

KẾ HOẠCH TRUYỀN THÔNG (tham chiếu):
{masterplan}

Trả lời CHỈ theo đúng định dạng JSON sau, không thêm chữ nào khác, không dùng markdown:
{{
  "passed": true hoặc false,
  "tone_issues": "mô tả vấn đề giọng văn nếu có, để chuỗi rỗng nếu ổn",
  "factual_issues": "mô tả thông tin bị bịa/nghi vấn nếu có, để chuỗi rỗng nếu ổn",
  "grammar_issues": "liệt kê lỗi chính tả/ngữ pháp cụ thể nếu có, để chuỗi rỗng nếu ổn",
  "length_issue": "nhận xét về độ dài nếu không phù hợp, để chuỗi rỗng nếu ổn",
  "overall_feedback": "nhận xét tổng quan ngắn gọn 1-2 câu"
}}"""
    user_msg = f"BÀI VIẾT CẦN KIỂM TRA (nền tảng: {req.platform}):\n\n{req.content}\n/no_think"

    client = InferenceClient(provider="nscale", api_key=HF_TOKEN)
    try:
        response = client.chat.completions.create(
            model="Qwen/Qwen3-32B",
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_msg},
            ],
            max_tokens=800,
        )
        raw = response.choices[0].message.content.strip()
        raw = re.sub(r'<think>.*?</think>', '', raw, flags=re.DOTALL).strip()
        if "```" in raw:
            raw = raw.split("```")[1]
            if raw.startswith("json"):
                raw = raw[4:]
        match = re.search(r'\{.*\}', raw, re.DOTALL)
        if match:
            raw = match.group(0)
        return json.loads(raw.strip())
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


# ---------------------------------------------------------------------------
# Exam bank (Reading/Listening, MCQ + short answer) — admin uploads, students take
# ---------------------------------------------------------------------------
import io
from docx import Document as DocxDocument


def require_admin(user):
    if user["role"] != "admin":
        raise HTTPException(status_code=403, detail="Chỉ admin mới được thực hiện thao tác này")


class ExamQuestionIn(BaseModel):
    question_number: int
    question_type: str  # "mcq" | "short_answer"
    question_text: str
    options: Optional[List[str]] = None
    correct_answer: str


@app.post("/exams/parse")
async def parse_exam(
    skill: str = Form(...),
    file: UploadFile = File(...),
    user=Depends(get_current_user),
):
    require_admin(user)
    if skill not in ("reading", "listening"):
        raise HTTPException(status_code=400, detail="Kỹ năng không hợp lệ")
    if not (file.filename or "").lower().endswith(".docx"):
        raise HTTPException(status_code=400, detail="Chỉ hỗ trợ file .docx")
    if not HF_TOKEN:
        raise HTTPException(status_code=500, detail="HF_TOKEN not configured")

    content = await file.read()
    try:
        doc = DocxDocument(io.BytesIO(content))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    paragraphs.append(" | ".join(cells))
        raw_text = "\n".join(paragraphs)
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Không đọc được file docx: {e}")

    if len(raw_text.strip()) < 30:
        raise HTTPException(status_code=400, detail="File docx không có nội dung hoặc quá ngắn")

    system_prompt = """Bạn là công cụ trích xuất đề thi IELTS từ văn bản thô (lấy ra từ file Word).
Nhiệm vụ: đọc văn bản, tách riêng phần bài đọc/bài nghe (passage/transcript) và danh sách câu hỏi kèm đáp án đúng.

Chỉ hỗ trợ 2 dạng câu hỏi:
- "mcq": trắc nghiệm nhiều lựa chọn, có field "options" là danh sách các lựa chọn (KHÔNG bao gồm ký tự A/B/C/D ở đầu mỗi lựa chọn).
- "short_answer": điền từ/câu trả lời ngắn, KHÔNG có field "options" (để null).

Nếu văn bản có đáp án/answer key tách riêng ở cuối, tự khớp đúng đáp án vào từng câu hỏi tương ứng theo đúng số thứ tự.

QUAN TRỌNG: Nếu KHÔNG tìm thấy đáp án cho 1 câu hỏi nào đó trong văn bản (đề không kèm đáp án), để "correct_answer" là chuỗi rỗng "". TUYỆT ĐỐI không tự bịa/đoán đáp án khi không có căn cứ trong văn bản gốc.

Trả lời CHỈ theo đúng định dạng JSON sau, không thêm chữ nào khác, không dùng markdown:
{
  "passage_text": "toàn bộ đoạn văn / transcript bài đọc hoặc bài nghe",
  "questions": [
    {
      "question_number": 1,
      "question_type": "mcq",
      "question_text": "...",
      "options": ["...", "...", "...", "..."],
      "correct_answer": "..."
    }
  ]
}"""

    client = InferenceClient(provider="nscale", api_key=HF_TOKEN)
    try:
        response = client.chat.completions.create(
            model="Qwen/Qwen3-32B",
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": raw_text[:12000] + "\n/no_think"},
            ],
            max_tokens=4000,
        )
        raw = response.choices[0].message.content.strip()
        raw = re.sub(r'<think>.*?</think>', '', raw, flags=re.DOTALL).strip()
        if "```" in raw:
            raw = raw.split("```")[1]
            if raw.startswith("json"):
                raw = raw[4:]
        match = re.search(r'\{.*\}', raw, re.DOTALL)
        if match:
            raw = match.group(0)
        return json.loads(raw.strip())
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Lỗi khi phân tích đề thi: {e}")


class ExamCreateRequest(BaseModel):
    title: str
    skill: str
    passage_text: str
    questions: List[ExamQuestionIn]


@app.post("/exams")
def create_exam(req: ExamCreateRequest, user=Depends(get_current_user)):
    require_admin(user)
    if req.skill not in ("reading", "listening"):
        raise HTTPException(status_code=400, detail="Kỹ năng không hợp lệ")
    if not req.questions:
        raise HTTPException(status_code=400, detail="Đề thi cần có ít nhất 1 câu hỏi")

    exam_row = {
        "title": req.title,
        "skill": req.skill,
        "passage_text": req.passage_text,
        "created_by": user["id"],
    }
    exam_res = supabase.table("exams").insert(exam_row).execute()
    exam_id = exam_res.data[0]["id"]

    question_rows = [{
        "exam_id": exam_id,
        "question_number": q.question_number,
        "question_type": q.question_type,
        "question_text": q.question_text,
        "options": q.options,
        "correct_answer": q.correct_answer,
    } for q in req.questions]
    supabase.table("exam_questions").insert(question_rows).execute()

    return {"status": "ok", "exam_id": exam_id}


@app.get("/exams")
def list_exams(skill: Optional[str] = None):
    require_supabase()
    q = supabase.table("exams").select("id,title,skill,created_at").order("created_at", desc=True)
    if skill:
        q = q.eq("skill", skill)
    return q.execute().data


@app.get("/exams/{exam_id}")
def get_exam(exam_id: str):
    require_supabase()
    exam = supabase.table("exams").select("*").eq("id", exam_id).single().execute()
    if not exam.data:
        raise HTTPException(status_code=404, detail="Không tìm thấy đề thi")
    questions = (
        supabase.table("exam_questions").select("*").eq("exam_id", exam_id)
        .order("question_number").execute().data
    )
    for q in questions:
        q.pop("correct_answer", None)  # an dap an, khong lo lo cho hoc sinh
    return {"exam": exam.data, "questions": questions}


class ExamSubmitRequest(BaseModel):
    answers: dict  # {question_id: cau_tra_loi}


@app.post("/exams/{exam_id}/submit")
def submit_exam(exam_id: str, req: ExamSubmitRequest):
    require_supabase()
    questions = (
        supabase.table("exam_questions").select("*").eq("exam_id", exam_id)
        .order("question_number").execute().data
    )
    if not questions:
        raise HTTPException(status_code=404, detail="Không tìm thấy đề thi")

    results = []
    correct_count = 0
    for q in questions:
        qid = q["id"]
        user_answer = (req.answers.get(qid) or "").strip()
        correct_answer = (q["correct_answer"] or "").strip()
        is_correct = user_answer.lower() == correct_answer.lower()
        if is_correct:
            correct_count += 1
        results.append({
            "question_id": qid,
            "question_number": q["question_number"],
            "your_answer": user_answer,
            "correct_answer": correct_answer,
            "is_correct": is_correct,
        })

    return {"score": correct_count, "total": len(questions), "results": results}
